"""
Weight Importer Utility
========================

Utility per importare pesi da vari formati (Excel, CSV, JSON, DOCX).
Supporta diversi formati di file per flessibilità massima.

Autore: Archaeological Classifier System
Data: Novembre 2025
"""

import pandas as pd
import json
import logging
from typing import Dict, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


def import_weights_from_excel(file_path: str,
                               artifact_id_col: str = None,
                               weight_col: str = None,
                               sheet_name: int | str = 0) -> Dict[str, float]:
    """
    Importa pesi da file Excel.

    Args:
        file_path: Path al file Excel (.xlsx, .xls)
        artifact_id_col: Nome colonna con ID artefatti (default: auto-detect)
        weight_col: Nome colonna con pesi (default: auto-detect)
        sheet_name: Nome o indice del foglio (default: primo foglio)

    Returns:
        Dict {artifact_id: weight_in_grams}

    Example:
        >>> weights = import_weights_from_excel('weights.xlsx')
        >>> print(weights['974'])
        387.0

    Formati Excel supportati:
        - Colonne: 'artifact_id', 'weight' (o 'peso', 'weight_g', etc.)
        - Colonne: 'ID', 'Weight (g)'
        - Colonne: 'Inventory', 'Mass'
    """
    try:
        logger.info(f"Importing weights from Excel: {file_path}")

        # Leggi Excel
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        logger.info(f"Excel loaded. Shape: {df.shape}, Columns: {list(df.columns)}")

        # Auto-detect colonne se non specificate
        if artifact_id_col is None:
            artifact_id_col = _detect_id_column(df)
            logger.info(f"Auto-detected ID column: {artifact_id_col}")

        if weight_col is None:
            weight_col = _detect_weight_column(df)
            logger.info(f"Auto-detected weight column: {weight_col}")

        if artifact_id_col not in df.columns:
            raise ValueError(f"Column '{artifact_id_col}' not found in Excel. "
                           f"Available: {list(df.columns)}")

        if weight_col not in df.columns:
            raise ValueError(f"Column '{weight_col}' not found in Excel. "
                           f"Available: {list(df.columns)}")

        # Estrai pesi
        weights = {}
        for idx, row in df.iterrows():
            artifact_id = str(row[artifact_id_col]).strip()
            weight_value = row[weight_col]

            # Skip se valori nulli
            if pd.isna(artifact_id) or pd.isna(weight_value):
                continue

            # Converti peso a float
            try:
                weight = float(weight_value)
                weights[artifact_id] = weight
            except (ValueError, TypeError) as e:
                logger.warning(f"Row {idx}: Could not convert weight '{weight_value}' to float: {e}")
                continue

        logger.info(f"Imported {len(weights)} weights from Excel")
        return weights

    except Exception as e:
        logger.error(f"Error importing from Excel: {e}", exc_info=True)
        raise


def import_weights_from_csv(file_path: str,
                            artifact_id_col: str = None,
                            weight_col: str = None,
                            delimiter: str = ',') -> Dict[str, float]:
    """
    Importa pesi da file CSV.

    Args:
        file_path: Path al file CSV
        artifact_id_col: Nome colonna ID (default: auto-detect)
        weight_col: Nome colonna peso (default: auto-detect)
        delimiter: Delimitatore CSV (default: ',')

    Returns:
        Dict {artifact_id: weight_in_grams}
    """
    try:
        logger.info(f"Importing weights from CSV: {file_path}")

        df = pd.read_csv(file_path, delimiter=delimiter)
        logger.info(f"CSV loaded. Shape: {df.shape}, Columns: {list(df.columns)}")

        # Auto-detect colonne
        if artifact_id_col is None:
            artifact_id_col = _detect_id_column(df)

        if weight_col is None:
            weight_col = _detect_weight_column(df)

        # Estrai pesi (stesso codice di Excel)
        weights = {}
        for idx, row in df.iterrows():
            artifact_id = str(row[artifact_id_col]).strip()
            weight_value = row[weight_col]

            if pd.isna(artifact_id) or pd.isna(weight_value):
                continue

            try:
                weight = float(weight_value)
                weights[artifact_id] = weight
            except (ValueError, TypeError) as e:
                logger.warning(f"Row {idx}: Could not convert weight '{weight_value}': {e}")
                continue

        logger.info(f"Imported {len(weights)} weights from CSV")
        return weights

    except Exception as e:
        logger.error(f"Error importing from CSV: {e}", exc_info=True)
        raise


def import_weights_from_json(file_path: str) -> Dict[str, float]:
    """
    Importa pesi da file JSON.

    Args:
        file_path: Path al file JSON

    Returns:
        Dict {artifact_id: weight_in_grams}

    Expected JSON format:
        {
            "974": 387.0,
            "942": 413.0,
            ...
        }
    """
    try:
        logger.info(f"Importing weights from JSON: {file_path}")

        with open(file_path, 'r') as f:
            data = json.load(f)

        # Converti tutto a float
        weights = {str(k): float(v) for k, v in data.items()}

        logger.info(f"Imported {len(weights)} weights from JSON")
        return weights

    except Exception as e:
        logger.error(f"Error importing from JSON: {e}", exc_info=True)
        raise


def import_weights_from_docx(file_path: str) -> Dict[str, float]:
    """
    Importa pesi da file DOCX (Note scansioni Artec formato Savignano).

    Cerca pattern come:
    - "974: 387g"
    - "Ascia 942 - 413 grammi"
    - Tabelle con ID e peso

    Args:
        file_path: Path al file DOCX

    Returns:
        Dict {artifact_id: weight_in_grams}
    """
    try:
        from docx import Document
        import re

        logger.info(f"Importing weights from DOCX: {file_path}")

        doc = Document(file_path)
        weights = {}

        # Pattern regex per trovare ID e peso
        # Esempi: "974: 387g", "Ascia 942 - 413 grammi", "ID: 974, Peso: 387g"
        patterns = [
            r'(\d+)\s*[:–-]\s*(\d+(?:\.\d+)?)\s*g',  # 974: 387g
            r'(\d+)\s*[:–-]\s*(\d+(?:\.\d+)?)\s*grammi',  # 974: 387 grammi
            r'ascia\s+(\d+)\s*[:–-]\s*(\d+(?:\.\d+)?)',  # Ascia 974 - 387
            r'ID:\s*(\d+).*?peso:\s*(\d+(?:\.\d+)?)',  # ID: 974, Peso: 387
        ]

        # Cerca in tutti i paragrafi
        for paragraph in doc.paragraphs:
            text = paragraph.text.lower()

            for pattern in patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    artifact_id = match.group(1)
                    weight_str = match.group(2)

                    try:
                        weight = float(weight_str)
                        weights[artifact_id] = weight
                        logger.debug(f"Found: {artifact_id} -> {weight}g in '{text[:50]}...'")
                    except ValueError:
                        continue

        # Cerca anche in tabelle
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]

                # Prova a trovare ID nella prima colonna, peso in seconda
                if len(cells) >= 2:
                    id_match = re.search(r'\d+', cells[0])
                    weight_match = re.search(r'(\d+(?:\.\d+)?)', cells[1])

                    if id_match and weight_match:
                        artifact_id = id_match.group(0)
                        weight_str = weight_match.group(1)

                        try:
                            weight = float(weight_str)
                            weights[artifact_id] = weight
                        except ValueError:
                            continue

        logger.info(f"Imported {len(weights)} weights from DOCX")

        if len(weights) == 0:
            logger.warning("No weights found in DOCX. Check format.")

        return weights

    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Error importing from DOCX: {e}", exc_info=True)
        raise


def import_weights_auto(file_path: str) -> Dict[str, float]:
    """
    Importa pesi rilevando automaticamente il formato dal file.

    Args:
        file_path: Path al file (Excel, CSV, JSON, o DOCX)

    Returns:
        Dict {artifact_id: weight_in_grams}

    Example:
        >>> weights = import_weights_auto('weights.xlsx')
        >>> weights = import_weights_auto('scan_notes.docx')
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext in ['.xlsx', '.xls']:
        return import_weights_from_excel(str(file_path))
    elif ext == '.csv':
        return import_weights_from_csv(str(file_path))
    elif ext == '.json':
        return import_weights_from_json(str(file_path))
    elif ext == '.docx':
        return import_weights_from_docx(str(file_path))
    else:
        raise ValueError(f"Unsupported file format: {ext}. "
                       f"Supported: .xlsx, .xls, .csv, .json, .docx")


# =============================================================================
# Helper Functions
# =============================================================================

def _detect_id_column(df: pd.DataFrame) -> str:
    """
    Auto-rileva colonna con ID artefatti.

    Cerca colonne con nomi comuni: 'id', 'artifact_id', 'inventory', etc.
    """
    id_candidates = [
        'artifact_id', 'id', 'artifact', 'artefact_id',
        'inventory', 'inventory_number', 'inv_num',
        'numero', 'num', 'n', 'code', 'codice'
    ]

    # Case-insensitive search
    for col in df.columns:
        col_lower = col.lower().strip()
        if col_lower in id_candidates:
            return col

    # Se non trovato, usa prima colonna
    logger.warning(f"Could not auto-detect ID column. Using first column: {df.columns[0]}")
    return df.columns[0]


def _detect_weight_column(df: pd.DataFrame) -> str:
    """
    Auto-rileva colonna con pesi.

    Cerca colonne con nomi comuni: 'weight', 'peso', 'mass', etc.
    """
    weight_candidates = [
        'weight', 'peso', 'mass', 'massa',
        'weight_g', 'weight_grams', 'peso_g',
        'grams', 'grammi', 'g', 'wt'
    ]

    # Case-insensitive search
    for col in df.columns:
        col_lower = col.lower().strip()
        if col_lower in weight_candidates:
            return col

        # Cerca anche substring
        for candidate in weight_candidates:
            if candidate in col_lower:
                return col

    # Se non trovato, prova seconda colonna (spesso: ID, Weight)
    if len(df.columns) >= 2:
        logger.warning(f"Could not auto-detect weight column. Using second column: {df.columns[1]}")
        return df.columns[1]

    raise ValueError("Could not auto-detect weight column and DataFrame has only 1 column. "
                   "Specify weight_col explicitly.")


# =============================================================================
# Example Usage
# =============================================================================

if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO)

    # Test Excel import
    print("Testing Excel import...")
    try:
        weights = import_weights_from_excel('test_weights.xlsx')
        print(f"Loaded {len(weights)} weights")
        print(f"Sample: {list(weights.items())[:3]}")
    except FileNotFoundError:
        print("test_weights.xlsx not found, skipping")

    # Test auto-detection
    print("\nTesting auto-detection...")
    test_file = "example_weights.xlsx"
    if Path(test_file).exists():
        weights = import_weights_auto(test_file)
        print(f"Loaded {len(weights)} weights")
